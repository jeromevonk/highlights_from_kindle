"""
# convert.py - Extract Kindle highlights and create Word documents
"""

import re
import os
from datetime import datetime
from docx import Document


def sanitize_filename(filename):
    """
    Remove or replace characters that are not allowed in filenames
    """
    # Replace problematic characters
    filename = re.sub(r'[<>:"/\\|?*]', '', filename)
    filename = filename.replace('...', '')
    filename = filename.strip()

    # Limit filename length
    if len(filename) > 100:
        filename = filename[:100]

    return filename


def extract_book_title(book_info):
    """
    Extract clean book title from the book info line
    """
    # Remove author info in parentheses and clean up
    title = re.sub(r'\([^)]*\)', '', book_info)
    title = title.strip()

    # Remove extra spaces
    title = re.sub(r'\s+', ' ', title)

    return title


def create_book_document(book_title, highlights, output_dir):
    """
    Create a Word document for a specific book with its highlights
    """
    # Create Word document
    doc = Document()

    # Add title
    doc.add_heading(f'Kindle Highlights - {book_title}', 0)

    # Add generation date
    date_para = doc.add_paragraph()
    date_para.add_run(f"Generated on: {datetime.now().strftime('%d/%m/%Y at %H:%M')}")

    # Add separator
    doc.add_paragraph("=" * 50)

    # Add highlights
    for i, highlight in enumerate(highlights, 1):
        # Add highlight number
        doc.add_heading(f'Highlight {i}', level=2)

        # Add metadata (page, position)
        if highlight['metadata']:
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(highlight['metadata']).italic = True

        # Add highlight text
        highlight_para = doc.add_paragraph()
        highlight_para.add_run(f'"{highlight["text"]}"')

        # Add separator between highlights
        if i < len(highlights):
            doc.add_paragraph("-" * 30)

    # Add summary at the end
    doc.add_page_break()
    doc.add_heading('Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Total highlights extracted: {len(highlights)}")

    # Create safe filename
    safe_title = sanitize_filename(book_title)
    filename = f"{safe_title}_Highlights.docx"
    filepath = os.path.join(output_dir, filename)

    # Save document
    try:
        doc.save(filepath)
        print(f"✓ Document saved: {filename} ({len(highlights)} highlights)")
        return True
    except Exception as e:
        print(f"✗ Error saving {filename}: {e}")
        return False


def extract_all_highlights(input_file_path, output_dir="highlights_output"):
    """
    Extract highlights from all books in Kindle highlights file
    and create separate Word documents for each book.
    """

    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Created directory: {output_dir}")

    # Read the input file
    try:
        with open(input_file_path, 'r', encoding='utf-8') as file:
            content = file.read()
    except FileNotFoundError:
        print(f"Error: File '{input_file_path}' not found.")
        return
    except Exception as e:
        print(f"Error reading file: {e}")
        return

    # Split content by separator lines
    highlights = content.split('==========')

    # Dictionary to store highlights by book
    books_highlights = {}

    for highlight in highlights:
        highlight = highlight.strip()
        if not highlight:
            continue

        lines = highlight.split('\n')

        # Extract book info, metadata, and highlight text
        book_info = lines[0] if lines else ""
        metadata = lines[1] if len(lines) > 1 else ""

        # Trim metadata to remove date information
        if metadata and "| Adicionado:" in metadata:
            metadata = metadata.split("| Adicionado:")[0].strip()

        highlight_text = '\n'.join(lines[2:]).strip() if len(lines) > 2 else ""

        # Only process if there's actual highlight text and book info
        if highlight_text and book_info:
            # Extract clean book title
            book_title = extract_book_title(book_info)

            # Initialize book in dictionary if not exists
            if book_title not in books_highlights:
                books_highlights[book_title] = []

            # Add highlight to book
            books_highlights[book_title].append({
                'book_info': book_info,
                'metadata': metadata,
                'text': highlight_text
            })

    if not books_highlights:
        print("No highlights found.")
        return

    # Create documents for each book
    print(f"\nFound {len(books_highlights)} books with highlights:")
    print("-" * 50)

    successful_docs = 0
    total_highlights = 0

    for book_title, highlights in books_highlights.items():
        if create_book_document(book_title, highlights, output_dir):
            successful_docs += 1
            total_highlights += len(highlights)

    # Print summary
    print("-" * 50)
    print("Summary:")
    print(f"• {successful_docs} documents created successfully")
    print(f"• {total_highlights} total highlights processed")
    print(f"• Files saved in: {output_dir}/")


def main():
    """
    Main function to run the highlight extraction process.
    """
    # File paths - adjust these as needed
    input_file = "My Clippings.txt"  # Your Kindle highlights file
    output_directory = "highlights"  # Directory for output files

    print("Extracting highlights from all books...")
    print("=" * 50)
    extract_all_highlights(input_file, output_directory)


if __name__ == "__main__":
    main()
